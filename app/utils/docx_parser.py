from pathlib import Path
from typing import Tuple
from loguru import logger


def parse_docx(file_path: str) -> Tuple[str, int]:
    """
    Parse a DOCX file and extract text content.
    Returns (text, page_count). Page count is estimated.
    """
    from docx import Document
    from docx.oxml.ns import qn

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc = Document(file_path)
        texts = []

        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    texts.append(row_text)

        full_text = "\n".join(texts)
        # Estimate page count: ~300 words per page
        word_count = len(full_text.split())
        estimated_pages = max(1, round(word_count / 300))

        return full_text, estimated_pages

    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_path}: {e}")
        raise ValueError(f"Could not parse DOCX: {file_path}") from e
